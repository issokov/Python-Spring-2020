import os
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError
from httplib2 import iri2uri
from urllib.request import urlopen
from urllib.error import HTTPError, URLError
from bs4 import BeautifulSoup
from error_handlers import ExecError
from collections import Iterable

class Executor:
    line_n = 0
    vars = {"_": None}
    loop_stack = []
    def __init__(self):
        self.actions = {
            "": (lambda: None, 0, 0),
            "set": (self.set, 1, 1000),
            "web-load-page": (self.load_page, 1, 2),
            "extract": (self.extract, 1, 3),
            "show": (self.show, 0, 1),
            "loop": (self.loop, 0, 2),
            "endloop": (self.end_loop, 0, 0),
            "create_dir": (self.create_dir, 1, 1),
            "web-save-file": (self.save_file, 2, 2),
            "doc-new": (self.doc_new, 0, 1),
            "doc-save": (self.doc_save, 1, 2),
            "doc": (self.doc, 1, 1000),
            "doc-album": (self.change_orient, 0, 1)
        }

    def exec_command(self, command, args):
        self.line_n += 1
        return self.actions[command][0](*args)

    def set(self, var_name, *expression):
        expression_ = " ".join(expression)
        try:
            self.vars[var_name] = eval(expression_, self.vars)
        except:
            return ExecError(ExecError.wrong_set, what=expression_)


    def load_data(self, iri):
        try:
            return urlopen(iri2uri(iri)).read(), None
        except HTTPError as e:
            return None, ExecError(ExecError.web_issue, f"Code: {e.code}\nReason: {e.reason}")
        except URLError as e:
            return None, ExecError(ExecError.web_issue, f"URL Error\nReason: {e.reason}")

    def load_page(self, var_url, output_var='_'):
        try:
            if var_url not in self.vars:
                return ExecError(ExecError.unknown_var, var_url)
            data, error = self.load_data(self.vars[var_url])
            if data:
                self.vars[output_var] = BeautifulSoup(data, "html.parser")
            else:
                return error
        except HTTPError as e:
            return ExecError(ExecError.web_issue, f"Code: {e.code}\nReason: {e.reason}")
        except URLError as e:
            return ExecError(ExecError.web_issue, f"URL Error\nReason: {e.reason}")

    def extract(self, what, target="_", output_var="_"):
        if target not in self.vars:
            return ExecError(ExecError.unknown_var, target)
        try:
            self.vars[output_var] = eval(f"self.vars[target].{what}")
        except:
            return ExecError(ExecError.extract_issue, f"Invalid extraction!\nQuery {target}.{what} isn't valid.")

    def show(self, var_name="_"):
        if var_name in self.vars:
            print(self.vars[var_name])
        else:
            return ExecError(ExecError.unknown_var, var_name)

    def loop(self, container_var_name="_", value_var_name="value"):
        if container_var_name in self.vars:
            container = self.vars[container_var_name]
            if isinstance(container, Iterable):
                if container:
                    self.loop_stack.append((value_var_name, (o for o in container), self.line_n))
                    try:
                        self.vars[value_var_name] = next(self.loop_stack[-1][1])
                    except StopIteration:
                        self.loop_stack.pop()
            else:
                return ExecError(ExecError.wrong_type,
                                 what=f"expected iterable, got: {type(self.vars[container_var_name])}")
        else:
            return ExecError(ExecError.unknown_var, container_var_name)

    def end_loop(self):
        if len(self.loop_stack):
            var_name, gen, start_line = self.loop_stack[-1]
            try:
                self.vars[var_name] = next(gen)
                self.line_n = start_line
            except StopIteration:
                self.loop_stack.pop()
        else:
            return ExecError(ExecError.loop_nf)

    def doc_new(self, var_name='_'):
        self.vars[var_name] = Document()

    def doc_save(self, path_var, var_name='_'):
        if var_name in self.vars:
            try:
                self.vars[var_name].save(self.vars[path_var])
            except:
                return ExecError(ExecError.wrong_type, f"Something goes wrong! Check {var_name} is it Document?")
        else:
            return ExecError(ExecError.unknown_var, f'{var_name}')

    def create_dir(self, path):
        try:
            os.mkdir(path)
        except FileExistsError:
            print(f'"System can`t create directory "{path}" because it`s already exists!')
        except:
            return ExecError(ExecError.mkdir_err, path)

    def save_file(self, iri_var, file_name_var="_"):
        if file_name_var in self.vars and iri_var in self.vars:
            data, error = self.load_data(self.vars[iri_var])
            if data:
                with open(self.vars[file_name_var], "wb") as file:
                    file.write(data)
            else:
                return error
        else:
            return ExecError(ExecError.unknown_var, file_name_var)

    def doc(self, *data):
        action = ' '.join(data[:-1])
        target = data[-1]
        if target not in self.vars:
            return ExecError(ExecError.unknown_var, target)
        try:
            locals().update(self.vars)
            exec(f"self.vars[target].{action}")
        except UnrecognizedImageError:
            exec("self.vars[target].add_paragraph('There should be a broken image')")
        except UnicodeDecodeError:
            exec("self.vars[target].add_paragraph('There should be a broken image')")
        except UnexpectedEndOfFileError:
            exec("self.vars[target].add_paragraph('There should be a broken image')")
        except:
            return ExecError(ExecError.doc_issue, f"Invalid operation!\nQuery {target}.{action} isn't valid.")

    def change_orient(self, doc_var="_"):
        if doc_var in self.vars:
            section = self.vars[doc_var].sections[-1]
            section.orientation = WD_ORIENTATION.LANDSCAPE
            # Bug with orientation
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        else:
            return ExecError(ExecError.unknown_var, doc_var)